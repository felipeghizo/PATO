from flask import Blueprint, render_template, request, flash, send_file
import logging
import os
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO
import re

# Configurações
UPLOAD_FOLDER = 'uploads/datasheetos'
ALLOWED_EXTENSIONS = {'pdf'}

# Configura logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

datasheetos_bp = Blueprint('datasheetos', __name__,
                         template_folder='templates',
                         static_folder='static')

@datasheetos_bp.route('/')
def datasheetos_home():
    return render_template('datasheetos/datasheetos.html')

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_specs(pdf_path):
    """Extrai especificações de um PDF"""
    specs = {}
    try:
        with open(pdf_path, 'rb') as file:
            reader = PdfReader(file)
            
            # Padrão para encontrar especificações (ajuste conforme necessário)
            pattern = re.compile(r'^([A-Z][^:]+):\s*(.+)$', re.MULTILINE)
            
            for page in reader.pages:
                text = page.extract_text()
                matches = pattern.finditer(text)
                for match in matches:
                    key = match.group(1).strip()
                    value = match.group(2).strip()
                    specs[key] = value
                    
    except Exception as e:
        logger.error(f"Erro ao extrair especificações: {str(e)}")
    
    return specs

def create_comparison_docx(dahua_specs, intelbras_specs):
    """Cria documento Word com especificações comuns"""
    doc = Document()
    doc.add_heading('Especificações Comuns', level=1)
    
    common_specs = set(dahua_specs.keys()) & set(intelbras_specs.keys())
    
    if not common_specs:
        doc.add_paragraph('Nenhuma especificação comum encontrada.')
        return doc
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Cabeçalhos da tabela
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Especificação'
    hdr_cells[1].text = 'Dahua'
    hdr_cells[2].text = 'Intelbras'
    
    # Adiciona linhas com especificações comuns
    for spec in sorted(common_specs):
        row_cells = table.add_row().cells
        row_cells[0].text = spec
        row_cells[1].text = dahua_specs[spec]
        row_cells[2].text = intelbras_specs[spec]
    
    return doc

@datasheetos_bp.route('/generate', methods=['POST'])
def generate_datasheet():
    try:
        # Verifica se os arquivos foram enviados
        if 'dahua_pdf' not in request.files or 'intelbras_pdf' not in request.files:
            flash('Por favor, envie ambos os arquivos PDF', 'error')
            return datasheetos_home()
        
        dahua_file = request.files['dahua_pdf']
        intelbras_file = request.files['intelbras_pdf']
        
        # Valida os arquivos
        if not (dahua_file and allowed_file(dahua_file.filename) and 
                intelbras_file and allowed_file(intelbras_file.filename)):
            flash('Apenas arquivos PDF são permitidos', 'error')
            return datasheetos_home()
        
        # Cria diretório de upload se não existir
        os.makedirs(UPLOAD_FOLDER, exist_ok=True)
        
        # Salva os arquivos temporariamente
        dahua_path = os.path.join(UPLOAD_FOLDER, secure_filename(dahua_file.filename))
        intelbras_path = os.path.join(UPLOAD_FOLDER, secure_filename(intelbras_file.filename))
        
        dahua_file.save(dahua_path)
        intelbras_file.save(intelbras_path)
        
        # Extrai especificações
        dahua_specs = extract_specs(dahua_path)
        intelbras_specs = extract_specs(intelbras_path)
        
        # Cria documento de comparação
        doc = create_comparison_docx(dahua_specs, intelbras_specs)
        
        # Remove arquivos temporários
        os.remove(dahua_path)
        os.remove(intelbras_path)
        
        # Retorna o documento Word para download
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return send_file(
            output,
            as_attachment=True,
            download_name='comparacao_especificacoes.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"Erro ao gerar datasheet: {str(e)}")
        flash('Ocorreu um erro ao processar os arquivos', 'error')
        return datasheetos_home()